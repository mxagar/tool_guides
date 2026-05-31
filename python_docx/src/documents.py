"""DOCX document generation helpers."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from datetime import timedelta
from decimal import Decimal
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from models import BudgetData, DocumentBaseData, ReceiptData


def money(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'))} EUR"


def _conversion_error_message(error: BaseException) -> str:
    if isinstance(error, subprocess.CalledProcessError):
        return error.stderr or error.stdout or str(error)
    return str(error)


class BaseDocument:
    """Common layout for company documents."""

    title = "Document"

    def __init__(self, data: DocumentBaseData):
        self.data = data
        self.document = Document()
        self._configure_styles()

    def build(self) -> Document:
        self._add_header()
        self._add_title()
        self._add_client_block()
        self._add_body()
        self._add_footer()
        return self.document

    def save_docx(self, path: Path | str) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.build()
        self.document.save(path)
        return path

    def save_pdf(self, docx_path: Path | str, pdf_path: Path | str | None = None) -> Path:
        docx_path = Path(docx_path)
        pdf_path = Path(pdf_path) if pdf_path else docx_path.with_suffix(".pdf")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        docx2pdf_error: BaseException | None = None
        try:
            from docx2pdf import convert

            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                return pdf_path
            docx2pdf_error = RuntimeError(f"docx2pdf did not create {pdf_path}")
        except (Exception, SystemExit) as error:
            docx2pdf_error = error

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError(
                "PDF export needs Microsoft Word via docx2pdf or LibreOffice/soffice "
                f"available on PATH. docx2pdf failed with: {docx2pdf_error}"
            ) from docx2pdf_error

        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(pdf_path.parent),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            generated_pdf = pdf_path.parent / docx_path.with_suffix(".pdf").name
            if not generated_pdf.exists():
                raise RuntimeError(f"LibreOffice did not create {generated_pdf}")
            if generated_pdf != pdf_path:
                generated_pdf.replace(pdf_path)
            return pdf_path
        except (subprocess.CalledProcessError, RuntimeError) as libreoffice_error:
            raise RuntimeError(
                "PDF export failed with docx2pdf and LibreOffice. "
                f"docx2pdf error: {docx2pdf_error}; "
                f"LibreOffice error: {_conversion_error_message(libreoffice_error)}"
            ) from libreoffice_error

    def export(self, output_dir: Path | str, filename_stem: str, pdf: bool = True) -> dict[str, Path | str]:
        output_dir = Path(output_dir)
        docx_path = self.save_docx(output_dir / f"{filename_stem}.docx")
        result: dict[str, Path | str] = {"docx": docx_path}
        if pdf:
            try:
                result["pdf"] = self.save_pdf(docx_path)
            except RuntimeError as error:
                result["pdf_error"] = str(error)
        return result

    def _configure_styles(self) -> None:
        styles = self.document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
        styles["Heading 1"].font.name = "Arial"
        styles["Heading 1"].font.size = Pt(18)

    def _add_header(self) -> None:
        company = self.data.company
        header = self.document.sections[0].header
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        left, right = table.rows[0].cells

        if company.logo_path and company.logo_path.exists():
            left.paragraphs[0].add_run().add_picture(str(company.logo_path), width=Inches(1.2))
        else:
            run = left.paragraphs[0].add_run(company.name)
            run.bold = True
            run.font.size = Pt(14)

        right_text = [
            company.name,
            f"Tax ID: {company.tax_id}",
            company.address.one_line(),
            f"{company.email} | {company.phone}",
        ]
        if company.website:
            right_text.append(company.website)
        right.paragraphs[0].text = "\n".join(right_text)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_title(self) -> None:
        heading = self.document.add_heading(self.title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = self.document.add_paragraph(f"Date: {self.data.document_date.isoformat()}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_client_block(self) -> None:
        client = self.data.client
        self.document.add_heading("Client", level=2)
        details = [
            client.full_name,
            f"Tax ID: {client.tax_id}",
            client.address.one_line(),
            client.email,
        ]
        if client.phone:
            details.append(client.phone)
        self.document.add_paragraph("\n".join(details))

    def _add_items_table(self) -> None:
        table = self.document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Description", "Unit price", "Quantity", "Unit", "Subtotal"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header

        for item in self.data.items:
            row = table.add_row().cells
            row[0].text = item.description
            row[1].text = money(item.unit_price)
            row[2].text = str(item.quantity.normalize())
            row[3].text = item.unit
            row[4].text = money(item.subtotal)

    def _add_totals(self) -> None:
        table = self.document.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        totals = [
            ("Net total", money(self.data.net_total)),
            (f"IVA ({self.data.iva_rate:.0%})", money(self.data.iva_total)),
            ("Gross total", money(self.data.gross_total)),
        ]
        for row, (label, value) in zip(table.rows, totals):
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_notes(self) -> None:
        if self.data.notes:
            self.document.add_heading("Notes", level=2)
            self.document.add_paragraph(self.data.notes)

    def _add_footer(self) -> None:
        footer = self.document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.text = f"{self.data.company.name} | {self.data.company.tax_id}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_body(self) -> None:
        raise NotImplementedError


class Budget(BaseDocument):
    title = "Budget"

    def __init__(self, data: BudgetData):
        super().__init__(data)
        self.data: BudgetData = data

    def _add_body(self) -> None:
        self.document.add_paragraph(f"Budget ID: {self.data.budget_id}")
        self.document.add_paragraph(f"Valid until: {self.data.valid_until.isoformat()}")
        self.document.add_heading("Products and Services", level=2)
        self._add_items_table()
        self.document.add_paragraph()
        self._add_totals()
        self._add_notes()


class Receipt(BaseDocument):
    title = "Receipt"

    def __init__(self, data: ReceiptData):
        super().__init__(data)
        self.data: ReceiptData = data

    def _add_body(self) -> None:
        due_date = self.data.document_date + timedelta(days=self.data.payment_due_days)
        self.document.add_paragraph(f"Receipt UUID: {self.data.receipt_uuid}")
        self.document.add_paragraph(f"Related budget ID: {self.data.budget_id}")
        self.document.add_paragraph(f"Payment due in {self.data.payment_due_days} days, by {due_date.isoformat()}.")
        self.document.add_paragraph(f"Bank transfer account: {self.data.bank_account}")
        self.document.add_heading("Products and Services", level=2)
        self._add_items_table()
        self.document.add_paragraph()
        self._add_totals()
        self._add_qr_code()
        self._add_notes()

    def _add_qr_code(self) -> None:
        qr_payload = (
            f"receipt_uuid={self.data.receipt_uuid};"
            f"budget_id={self.data.budget_id};"
            f"gross_total={money(self.data.gross_total)}"
        )
        image = qrcode.make(qr_payload)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = Path(tmp.name)
            image.save(tmp_path)
        try:
            self.document.add_heading("Payment QR", level=2)
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(tmp_path), width=Inches(1.5))
        finally:
            tmp_path.unlink(missing_ok=True)
